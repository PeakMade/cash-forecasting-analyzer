"""
Word Document Generator Service
Creates professional DOCX reports from analysis recommendations
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from typing import Dict, Any
import logging

logger = logging.getLogger(__name__)


class WordDocumentGenerator:
    """Generate Word documents from recommendation data"""
    
    def __init__(self):
        # PeakMade/Redpoint brand colors
        self.brand_blue = RGBColor(65, 105, 225)  # Royal Blue
        self.brand_orange = RGBColor(255, 140, 0)  # Dark Orange
        self.dark_gray = RGBColor(51, 51, 51)
        self.light_gray = RGBColor(128, 128, 128)
        
    def generate_document(self, recommendation: Dict[str, Any], output_path: str) -> str:
        """
        Generate Word document from recommendation data
        
        Args:
            recommendation: Complete recommendation dictionary from RecommendationEngine
            output_path: Path where DOCX file should be saved
            
        Returns:
            Path to generated DOCX file
        """
        logger.info(f"Generating Word document for {recommendation.get('property_name', 'Unknown')}")
        
        try:
            doc = Document()
            
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Title page
            self._add_title_section(doc, recommendation)
            
            # Executive Summary
            self._add_executive_summary(doc, recommendation)
            
            # Detailed Analysis sections
            self._add_detailed_analysis(doc, recommendation)
            
            # Save document
            doc.save(output_path)
            logger.info(f"Word document saved to {output_path}")
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating Word document: {str(e)}")
            raise
    
    def _add_title_section(self, doc: Document, recommendation: Dict[str, Any]):
        """Add title and property information"""
        # Property name as title
        title = doc.add_heading(recommendation.get('property_name', 'Unknown Property'), level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = self.dark_gray
        title_run.font.size = Pt(28)
        
        # Analysis period
        period = doc.add_paragraph(f"{recommendation.get('analysis_month', 'Unknown')} Analysis")
        period.alignment = WD_ALIGN_PARAGRAPH.CENTER
        period_run = period.runs[0]
        period_run.font.size = Pt(14)
        period_run.font.color.rgb = self.light_gray
        period_run.font.italic = True
        
        doc.add_paragraph()  # Spacing
        
        # Executive Decision
        decision = recommendation.get('decision', 'DO_NOTHING')
        amount = recommendation.get('amount')
        
        decision_para = doc.add_heading('EXECUTIVE DECISION', level=2)
        decision_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        decision_text = doc.add_paragraph()
        decision_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        decision_display = decision.replace('_', ' ')
        decision_run = decision_text.add_run(f"{decision_display}")
        decision_run.font.size = Pt(24)
        decision_run.bold = True
        
        # Color code by decision type
        if decision == 'CONTRIBUTE':
            decision_run.font.color.rgb = RGBColor(220, 53, 69)  # Red
        elif decision == 'DISTRIBUTE':
            decision_run.font.color.rgb = RGBColor(40, 167, 69)  # Green
        else:
            decision_run.font.color.rgb = RGBColor(23, 162, 184)  # Blue
        
        if amount:
            amount_para = doc.add_paragraph()
            amount_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            amount_run = amount_para.add_run(f"${amount:,.0f}")  # Removed decimal places since rounding to $10K
            amount_run.font.size = Pt(20)
            amount_run.bold = True
            amount_run.font.color.rgb = self.dark_gray
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc: Document, recommendation: Dict[str, Any]):
        """Add executive summary bullets"""
        doc.add_heading('Executive Summary', level=1)
        
        executive_summary = recommendation.get('executive_summary', [])
        
        for bullet in executive_summary:
            p = doc.add_paragraph(bullet, style='List Bullet')
            p.paragraph_format.space_after = Pt(6)
        
        doc.add_paragraph()  # Spacing
    
    def _add_detailed_analysis(self, doc: Document, recommendation: Dict[str, Any]):
        """Add all detailed analysis sections"""
        detailed = recommendation.get('detailed_rationale', {})
        
        # Decision Rationale
        if detailed.get('decision_rationale'):
            doc.add_heading('Decision Rationale', level=1)
            doc.add_paragraph(detailed['decision_rationale'])
            doc.add_paragraph()
        
        # Cash Forecast Analysis
        if detailed.get('cash_forecast_analysis'):
            doc.add_heading('Cash Forecast Analysis', level=1)
            doc.add_paragraph(detailed['cash_forecast_analysis'])
            doc.add_paragraph()
        
        # Income Statement Analysis
        if detailed.get('income_statement_analysis'):
            doc.add_heading('Income Statement Analysis', level=1)
            doc.add_paragraph(detailed['income_statement_analysis'])
            doc.add_paragraph()
        
        # Balance Sheet Analysis
        if detailed.get('balance_sheet_analysis'):
            doc.add_heading('Balance Sheet Analysis', level=1)
            doc.add_paragraph(detailed['balance_sheet_analysis'])
            doc.add_paragraph()
        
        # Economic & Market Context
        if detailed.get('economic_context'):
            doc.add_heading('Economic & Market Context', level=1)
            doc.add_paragraph(detailed['economic_context'])
            doc.add_paragraph()
        
        # Financial Metrics
        if detailed.get('financial_metrics'):
            doc.add_heading('Financial Metrics', level=1)
            doc.add_paragraph(detailed['financial_metrics'])
            doc.add_paragraph()
        
        # Cash Flow Analysis
        if detailed.get('cash_flow_analysis'):
            doc.add_heading('Cash Flow Analysis', level=1)
            doc.add_paragraph(detailed['cash_flow_analysis'])
            doc.add_paragraph()
        
        # Market Context
        if detailed.get('market_context'):
            doc.add_heading('Market & Economic Context', level=1)
            doc.add_paragraph(detailed['market_context'])
            doc.add_paragraph()
        
        # Historical Performance
        if detailed.get('historical_performance'):
            doc.add_heading('Historical Performance', level=1)
            doc.add_paragraph(detailed['historical_performance'])
            doc.add_paragraph()
        
        # Risk Assessment
        if detailed.get('risk_assessment'):
            doc.add_heading('Risk Assessment', level=1)
            doc.add_paragraph(detailed['risk_assessment'])
            doc.add_paragraph()
        
        # Recommendations
        if detailed.get('recommendations'):
            doc.add_heading('Recommendations', level=1)
            doc.add_paragraph(detailed['recommendations'])
